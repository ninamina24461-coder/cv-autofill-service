from __future__ import annotations

import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from .utils import clean_text

# Lazy loading for heavy dependencies
_pdfplumber = None
_docx = None

def get_pdfplumber():
    global _pdfplumber
    if _pdfplumber is None:
        import pdfplumber
        _pdfplumber = pdfplumber
    return _pdfplumber

def get_docx():
    global _docx
    if _docx is None:
        from docx import Document
        _docx = Document
    return _docx


@dataclass
class TextExtraction:
    text: str
    mime_type: Optional[str]
    warnings: List[str]


def extract_text(filename: str, content: bytes, mime_type: Optional[str] = None) -> TextExtraction:
    suffix = Path(filename).suffix.lower()
    warnings: List[str] = []

    if suffix == ".pdf" or mime_type == "application/pdf":
        text, pdf_warnings = _from_pdf(content)
        warnings.extend(pdf_warnings)
        return TextExtraction(text=text, mime_type="application/pdf", warnings=warnings)

    if suffix == ".docx" or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return TextExtraction(text=_from_docx(content), mime_type=mime_type, warnings=warnings)

    if suffix in {".txt", ".md", ".rtf"} or (mime_type or "").startswith("text/"):
        try:
            return TextExtraction(text=content.decode("utf-8", errors="ignore"), mime_type="text/plain", warnings=warnings)
        except Exception as exc:
            return TextExtraction(text="", mime_type="text/plain", warnings=[f"text_decode_failed:{exc}"])

    warnings.append(f"unsupported_file_type:{suffix or mime_type or 'unknown'}")
    return TextExtraction(text="", mime_type=mime_type, warnings=warnings)


def _from_pdf(content: bytes):
    warnings: List[str] = []
    chunks: List[str] = []
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(content)
        tmp.flush()
        tmp_path = tmp.name
    
    try:
        pdfplumber = get_pdfplumber()
        with pdfplumber.open(tmp_path) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                if txt.strip():
                    chunks.append(txt)
    finally:
        import os
        os.unlink(tmp_path)
    
    text = "\n".join(chunks).strip()
    if not text:
        warnings.append("pdf_text_empty")
    return text, warnings


def _from_docx(content: bytes) -> str:
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(content)
            tmp.flush()
            tmp_path = tmp.name
        
        Document = get_docx()
        doc = Document(tmp_path)
        parts: List[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts).strip()
    finally:
        if tmp_path:
            import os
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
